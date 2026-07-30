"""Generate the full editable Word report (faithful reproduction + low-data win)."""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import re

BASE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUT = os.path.join(BASE, "report", "final_report_v2.docx")
NAVY = RGBColor(0x1a, 0x3c, 0x6e)

doc = Document()
doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(10.5)


def runs(p, text):
    for chunk in re.split(r"(\*\*.*?\*\*|\*.*?\*)", text):
        if not chunk:
            continue
        if chunk.startswith("**") and chunk.endswith("**"):
            p.add_run(chunk[2:-2]).bold = True
        elif chunk.startswith("*") and chunk.endswith("*"):
            p.add_run(chunk[1:-1]).italic = True
        else:
            p.add_run(chunk)


def para(text, size=10.5, after=6, align=None):
    p = doc.add_paragraph(); runs(p, text)
    for r in p.runs: r.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(after)
    if align: p.alignment = align
    return p


def heading(text, size=13):
    p = doc.add_paragraph(); r = p.add_run(text)
    r.bold = True; r.font.size = Pt(size); r.font.color.rgb = NAVY
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(3)


def caption(text):
    p = doc.add_paragraph(); r = p.add_run(text)
    r.italic = True; r.font.size = Pt(8); r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(8)


def figure(rel, width_cm=13):
    path = os.path.join(BASE, rel)
    if os.path.exists(path):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(path, width=Cm(width_cm))


def table(rows, widths):
    t = doc.add_table(rows=0, cols=len(rows[0])); t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(rows):
        cells = t.add_row().cells
        for j, val in enumerate(row):
            cells[j].text = str(val)
            for para_ in cells[j].paragraphs:
                for r in para_.runs:
                    r.font.size = Pt(9)
                    if i == 0: r.bold = True
    return t


# ---------------- Title ----------------
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("Context-Gated Channel Attention:\nExtending the Attention U-Net with Channel Selection")
r.bold = True; r.font.size = Pt(16); r.font.color.rgb = NAVY
para("Medical Images Processing with Deep Learning (336033) — Final Project",
     size=10.5, after=4, align=WD_ALIGN_PARAGRAPH.CENTER)
para("**Aviv Niemann** (ID: __________)    ·    **______________** (ID: __________)",
     after=8, align=WD_ALIGN_PARAGRAPH.CENTER)
para("**Selected paper:** O. Oktay et al., \"Attention U-Net: Learning Where to Look for the "
     "Pancreas,\" MIDL 2018 (arXiv:1804.03999)", size=9.5, after=4, align=WD_ALIGN_PARAGRAPH.CENTER)
para("**Proposed extension:** add channel attention to the attention gate's skip-connection "
     "gating — a novel *context-gated* channel gate (conditioned on the coarse gating signal), "
     "compared against a CBAM-style channel-attention variant.", size=9.5, after=10,
     align=WD_ALIGN_PARAGRAPH.CENTER)

# ---------------- 1. Introduction ----------------
heading("1. Introduction")
para("**Summary of the paper.** The Attention U-Net augments the standard U-Net with "
     "*attention gates* (AGs) on the skip connections. Each gate combines the fine-scale "
     "encoder features with a coarser-scale *gating signal g* from the decoder to produce a "
     "per-pixel spatial attention map that suppresses background responses before they reach "
     "the decoder — letting the network learn *where to look* without an external localization "
     "step. AGs are lightweight, trained end-to-end, and improve segmentation of small, "
     "variable organs, with the benefit most pronounced when training data is limited.")
para("**Problem and relevance.** Automated organ segmentation in CT is clinically valuable but "
     "hard for small, low-contrast, shape-variable structures. Attention gates fold organ "
     "localization into a single end-to-end model, improving accuracy at negligible cost and "
     "adding a degree of interpretability.")
para("**Motivation.** The attention gate is purely *spatial*: it decides *where* to look but "
     "applies the same scalar to every feature channel — it never decides *which channels* are "
     "relevant. Since different channels encode different cues (edges, texture, semantics), we "
     "hypothesized that adding **channel** selection to the gate could improve it, and we chose "
     "this paper because that gap is a clear, well-motivated place to extend a widely used method.")

# ---------------- 2. Proposed Extension ----------------
heading("2. Proposed Extension")
para("**Our idea — a context-gated channel gate (Hybrid).** We add a second gating branch that "
     "outputs a per-channel weight, applied alongside the original spatial coefficient. Crucially, "
     "this channel weight is computed from *both* the local features and the same coarse gating "
     "signal *g* the spatial gate uses — extending the paper's top-down \"where to look\" "
     "principle from spatial locations to feature channels. The gate is initialized to pass all "
     "channels so it cannot suppress useful features before learning.")
para("**Reasoning and hypothesis.** The original AG collapses all channel information into one "
     "spatial scalar. We hypothesized that a context-conditioned channel gate would improve the "
     "precision/recall trade-off, especially for small, low-contrast structures. Our objective "
     "was to *improve on the paper's model* — the spatial Attention U-Net — not merely the plain "
     "U-Net.")
para("**Alternatives considered, and how the study evolved.** When our context-gated Hybrid did "
     "not improve over the paper's model at full data (Section 5), we drew on the literature and "
     "adopted a second, established channel-attention design as a comparison: **CBAM** "
     "(Woo et al., ECCV 2018), which applies channel-then-spatial self-attention. Unlike our "
     "gate, CBAM's channel attention uses only the feature map's own statistics (it ignores *g*). "
     "We combined it with the attention gate in the manner of ASCU-Net (2021). Finally, because "
     "the original paper reports that attention helps most with scarce data, we designed a "
     "**low-data experiment** to test whether either channel-attention design improves the "
     "paper's model when training data is limited.")

# ---------------- 3. Methodology ----------------
heading("3. Methodology")
para("**Models.** All four variants are built from one configurable 2D U-Net so the *only* "
     "difference between them is the skip-connection gate: **U-Net** (no gate); **Attention "
     "U-Net** — the paper's model (spatial grid gate, ported from the authors' code); **AG+CBAM** "
     "(spatial gate + CBAM channel & spatial self-attention); and **Hybrid — ours** (spatial gate "
     "+ context-gated channel gate). All use **deep supervision** (auxiliary heads at each decoder "
     "scale), following the paper.")
para("**Dataset and preprocessing.** Medical Segmentation Decathlon **Task09_Spleen** (41 labeled "
     "3D abdominal CT volumes). Volumes are windowed to a soft-tissue range, sliced to 2D axial "
     "slices, resized to 256×256, and per-slice **z-score normalized**. Patients are split at the "
     "**patient level** (never per slice) to avoid leakage.")
para("**Training.** Adam (lr 3×10⁻⁴), batch 8, up to 60 epochs with early stopping on validation "
     "Dice, gradient clipping, and paper-style augmentation (affine rotation/scale + flips). "
     "**Loss:** the original paper uses Dice loss alone (on 3D pancreas volumes). On our 2D spleen "
     "slices, however, the spleen occupies only a tiny fraction of each image — and many slices "
     "contain no spleen at all — so under pure Dice loss the model can minimize its error by "
     "simply predicting \"all background\"; once it does, Dice's gradient becomes vanishingly small "
     "and it cannot recover, collapsing to an empty prediction. We therefore add a **focal** term, "
     "which keeps a strong learning signal on the few foreground pixels and eliminates the collapse "
     "(applied identically to all variants). This instability did not arise in the original paper "
     "because full 3D volumes contain the whole organ per sample, making the class imbalance far "
     "less extreme. "
     "**Evaluation:** the full-data study uses **5-fold cross-validation**; the low-data study "
     "uses a fixed held-out test set with multiple seeds (the standard design for a data-size "
     "sweep, and how the paper studied training-size effects). We report Dice, precision, recall, "
     "and paired Wilcoxon signed-rank tests on per-slice Dice.")
para("**Tools.** PyTorch, MONAI, nibabel, scikit-image, SciPy, matplotlib. Experiments ran on an "
     "NVIDIA L40 GPU. (We first attempted a full 3D reproduction; a naïve 3D pipeline failed to "
     "converge — a stock reference U-Net also underperformed through it — so we conducted the "
     "study in 2D, where the baseline reproduces expected performance.)")

# ---------------- 4. (placeholder heading merged into results) ----------------

# ---------------- 5. Results ----------------
heading("4. Results and Analysis")
para("**Full data: our extension does not beat the paper's model.** Under 5-fold "
     "cross-validation, all four variants score similarly (per-fold Dice below). Our Hybrid is "
     "*significantly worse* than the paper's Attention U-Net (paired Wilcoxon p<0.001), and the "
     "CBAM variant's small edge is **within the fold-to-fold standard deviation and not "
     "significant** (p=0.27). In short, with abundant data on this (relatively easy) organ, "
     "adding channel attention gives no reliable improvement — there is little headroom above the "
     "already-strong baseline.")
table([["Variant", "Dice (mean ± std)", "Precision", "Recall"],
       ["U-Net", "0.856 ± 0.061", "0.906", "0.885"],
       ["Attention U-Net (paper)", "0.869 ± 0.048", "0.908", "0.895"],
       ["AG + CBAM", "0.872 ± 0.034", "0.908", "0.892"],
       ["Hybrid (ours)", "0.852 ± 0.079", "0.903", "0.878"]],
      None)
caption("Table 1. Full-data 5-fold cross-validation on spleen. Differences are small and overlap "
        "within std; the CBAM edge over the paper's model is not significant (p=0.27).")
figure("results_faithful2d/fig_kfold_box.png", width_cm=11)
caption("Figure 1. Per-fold Dice (full data). The four variants overlap almost entirely; "
        "fold-to-fold variation exceeds any difference between variants.")

para("**Low data: channel attention significantly improves the paper's model.** Motivated by the "
     "paper's claim that attention helps most with scarce data, we trained each variant on only "
     "**8 patients** (fixed test set, 8 seeds; ~2,000 paired test slices). Here both "
     "channel-attention designs **significantly outperform the paper's Attention U-Net**: AG+CBAM "
     "by +0.084 Dice and our Hybrid by +0.031 Dice (paired Wilcoxon p<0.0001 for both). AG+CBAM "
     "is strongest; our Hybrid — which lost at full data — now clearly beats the paper's model.")
table([["Variant", "Dice (n=8, 8 seeds)", "vs paper's model"],
       ["U-Net", "0.784", "—"],
       ["Attention U-Net (paper)", "0.793", "baseline"],
       ["AG + CBAM", "0.877", "+0.084  (p<0.0001)"],
       ["Hybrid (ours)", "0.824", "+0.031  (p<0.0001)"]],
      None)
caption("Table 2. Low-data regime (8 training patients). Both channel-attention variants "
        "significantly beat the paper's model; effect sizes +0.03 to +0.08 Dice.")
figure("results_lowdata_focus/fig_lowdata_n8.png", width_cm=11)
caption("Figure 2. Low-data (n=8) mean Dice. Channel attention (AG+CBAM, and our Hybrid) "
        "significantly exceeds the paper's spatial-only Attention U-Net.")

para("**Qualitative comparison.** Figure 3 shows predictions on fixed test slices for the "
     "low-data models. The U-Net and the paper's Attention U-Net produce frequent misses and "
     "stray false positives (e.g. the Attention U-Net scores 0.00 on one slice), whereas AG+CBAM "
     "and our Hybrid track the spleen far more cleanly — the visual counterpart of the Dice gains. "
     "At full data the four models look nearly identical (consistent with Table 1).")
figure("results_grids/fig_grid_lowdata.png", width_cm=15)
caption("Figure 3. Low-data (n=8) predictions (red) vs. ground truth (blue). Channel-attention "
        "variants (right two columns) are visibly cleaner than the baseline and the paper's model.")

para("**Failure cases and challenges.** Training was initially unstable: pure Dice loss "
     "collapsed to all-background predictions on small-foreground slices (fixed with the Dice+Focal "
     "loss and open-gate initialization), and a 3D pipeline failed to converge (diagnosed via a "
     "stock reference U-Net). The gated variants also show higher seed-to-seed variance than the "
     "plain U-Net, and per-slice significance p-values are anti-conservative (slices from one "
     "patient are correlated) — so we emphasize effect sizes over p-values throughout.")

# ---------------- 6. Conclusion ----------------
heading("5. Conclusion")
para("**Findings.** We proposed a context-gated channel-attention extension to the Attention "
     "U-Net and evaluated it rigorously. At full data it does *not* improve the paper's model "
     "(and a CBAM variant's edge is within noise). However, in the **low-data regime** both "
     "channel-attention designs — including our Hybrid — **significantly improve the paper's "
     "model** (Hybrid +0.031, CBAM +0.084 Dice, p<0.0001). This confirms and extends the paper's "
     "own thesis: attention's benefit concentrates when data is scarce, and *channel* selection "
     "adds value on top of spatial attention precisely in that regime.")
para("**Limitations.** A 2D proof-of-concept on a single, relatively easy organ; the required "
     "Dice+Focal loss adaptation; higher variance for the gated variants; a fixed-test (not "
     "k-fold) low-data protocol; and CBAM outperforming our Hybrid (the simpler channel attention "
     "is stronger). Per-slice significance overstates certainty due to correlated slices.")
para("**Future work.** Extend to full 3D volumes (where the paper operated); test harder organs "
     "(e.g. pancreas) and thin/branching structures; run k-fold at each training size and more "
     "seeds to firm up the low-data comparison; and investigate why context-conditioning helps "
     "less than simple channel recalibration, to improve the Hybrid design.")

doc.save(OUT)
print("wrote", OUT)
