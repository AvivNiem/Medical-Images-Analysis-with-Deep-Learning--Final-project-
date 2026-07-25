"""make_report_docx.py — editable Word (.docx) version of the final report.
Embeds the real result numbers, the results table, and the figures from results/."""
import os
import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

HERE = os.path.dirname(os.path.abspath(__file__))
RESULTS = os.path.join(os.path.dirname(HERE), "results")
OUT = os.path.join(HERE, "final_report.docx")
NAVY = RGBColor(0x1a, 0x3c, 0x6e)

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10.5)


def add_runs(par, text):
    """Render simple **bold** / *italic* markup into runs."""
    for chunk in re.split(r"(\*\*.*?\*\*|\*.*?\*)", text):
        if not chunk:
            continue
        if chunk.startswith("**") and chunk.endswith("**"):
            par.add_run(chunk[2:-2]).bold = True
        elif chunk.startswith("*") and chunk.endswith("*"):
            par.add_run(chunk[1:-1]).italic = True
        else:
            par.add_run(chunk)


def para(text, size=10.5, space_after=6, align=None):
    p = doc.add_paragraph()
    add_runs(p, text)
    for r in p.runs:
        r.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(space_after)
    if align:
        p.alignment = align
    return p


def heading(text, size=13):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)
    r.font.color.rgb = NAVY
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    return p


def caption(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)


def figure(name, width_cm=13):
    path = os.path.join(RESULTS, name)
    if os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(path, width=Cm(width_cm))


# ---- Title block ----
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("Context-Gated Channel Attention:\nExtending Attention U-Net with Top-Down Channel Selection")
r.bold = True
r.font.size = Pt(16)
r.font.color.rgb = NAVY

para("Medical Images Processing with Deep Learning (336033) — Final Project",
     size=10.5, space_after=4, align=WD_ALIGN_PARAGRAPH.CENTER)
para("**Aviv Niemann** (ID: __________)    ·    **______________** (ID: __________)",
     size=10.5, space_after=8, align=WD_ALIGN_PARAGRAPH.CENTER)
para("**Selected paper:** O. Oktay et al., \"Attention U-Net: Learning Where to Look for the "
     "Pancreas,\" MIDL 2018 (arXiv:1804.03999)", size=9.5, space_after=4,
     align=WD_ALIGN_PARAGRAPH.CENTER)
para("**Proposed extension:** a context-gated channel-attention branch added to the attention "
     "gate, conditioned on the same coarse-scale gating signal the original spatial gate uses — "
     "extending the paper's top-down \"where to look\" mechanism from spatial locations to feature "
     "channels.", size=9.5, space_after=10, align=WD_ALIGN_PARAGRAPH.CENTER)

# ---- 1. Introduction ----
heading("1. Introduction")
para("**Summary of the paper.** The Attention U-Net augments the standard U-Net with *attention "
     "gates* (AGs) on the skip connections. Each gate takes the fine-scale encoder features x that "
     "would be concatenated into the decoder together with a coarser-scale *gating signal g* from "
     "the deeper decoder path, and produces a per-pixel spatial attention coefficient α ∈ [0,1]. "
     "Multiplying x by α suppresses background responses before they reach the decoder, so the "
     "network learns to focus on the target organ without an external localization module. The "
     "gates are additive, trained end-to-end, and add few parameters. On abdominal CT the authors "
     "show AGs improve pancreas segmentation — notably increasing recall — with the benefit most "
     "pronounced when training data is scarce.")
para("**Problem and relevance.** Automated organ segmentation in CT is clinically valuable but "
     "hard for small, low-contrast, shape-variable structures. Prior state of the art relied on "
     "multi-stage cascades (localize, then segment), which are redundant and complex. The attention "
     "gate folds localization into a single end-to-end model for negligible cost, and its attention "
     "maps offer a degree of interpretability.")
para("**Motivation.** The gate is elegant and widely used, but has a clear limitation: its "
     "attention is purely *spatial* — it decides *where* to look but applies the same scalar across "
     "all feature channels, never deciding *which* channels are relevant. Channel-attention modules "
     "(SE-Net, CBAM) address channel selection, but via pure self-attention over the feature map's "
     "own statistics, discarding the paper's core idea of top-down contextual gating. This gap is a "
     "natural place to extend the method.")

# ---- 2. Proposed Extension ----
heading("2. Proposed Extension")
para("**Description.** We add a second gating branch that produces a per-channel weight β ∈ "
     "[0,1]^C applied alongside the spatial coefficient, so the gated skip output becomes "
     "W(x · α · β). Crucially β is computed from *both* the local features x and the coarse gating "
     "signal g, mirroring the spatial gate: β = σ(W_out · ReLU(W_x·GAP(x) + W_g·GAP(g))), where GAP "
     "is global average pooling. The channel gate is thus *context-gated*: the coarse decoder "
     "context helps decide which channels matter, just as it already decides which locations "
     "matter. It is initialized to pass all channels (β ≈ 1) so it cannot suppress useful features "
     "before learning.")
para("**Gap addressed and hypothesis.** The original AG collapses channel information into a single "
     "spatial scalar. We hypothesize that different channels carry different, location-dependent "
     "relevance — especially for small, low-contrast organs — and that letting top-down context "
     "modulate channels improves the precision/recall trade-off. We further hypothesize the "
     "advantage over both the baseline and a naïve channel-attention bolt-on *grows as training "
     "data shrinks*, the regime the original paper highlights as most favorable to attention.")
para("**Alternatives considered.** (i) A standard CBAM block on the AG — its channel attention is "
     "pure self-attention and ignores g; we retain it as a *control* to isolate the value of "
     "context-conditioning. (ii) Squeeze-and-Excitation in the encoder — same objection, not tied "
     "to skip gating. (iii) Residual/highway connections around the gate — the original authors "
     "reported no benefit, so we did not pursue it.")

# ---- 3. Methodology ----
heading("3. Methodology")
para("**Models.** All four variants are built from one configurable 2D U-Net so the only "
     "difference is the gating module: **U-Net** (no gating); **Attention U-Net** (original spatial "
     "gate α(x,g)); **AG+CBAM** (spatial gate + CBAM channel/spatial self-attention, the control); "
     "and **Hybrid, ours** (spatial gate α(x,g) and context-gated channel gate β(x,g)). Following "
     "the paper, the shallowest skip is left ungated; the spatial gate reproduces the authors' "
     "reference implementation.")
para("**Data and preprocessing.** Medical Segmentation Decathlon Task09_Spleen — 41 labeled 3D "
     "abdominal CT volumes with binary spleen masks, chosen to stay in the paper's CT small-organ "
     "domain while remaining a feasible proof-of-concept. Volumes are windowed to a soft-tissue HU "
     "range, normalized, sliced to 2D axial slices and resized to 256×256; all spleen slices plus a "
     "sample of background slices are kept. Patients are split train/val/test at the **patient "
     "level** (never per slice) to prevent leakage; the same test set is used for every variant and "
     "training size.")
para("**Training.** Adam (lr 3×10⁻⁴), batch 8, up to 60 epochs with early stopping on validation "
     "Dice, gradient clipping, light augmentation, 5 random seeds; we report the mean and a paired "
     "Wilcoxon signed-rank test on per-slice Dice. **Loss — a stability finding:** the paper uses "
     "Dice loss for imbalance robustness, but on our small-foreground 2D slices pure Dice training "
     "was stochastically unstable, intermittently collapsing to an all-background prediction from "
     "which Dice's vanishing gradient could not recover; plain BCE+Dice was worse under heavy "
     "imbalance. We adopted a **Dice+Focal** compound loss — Dice optimizes overlap, focal "
     "down-weights easy background and preserves a strong foreground gradient — which eliminated the "
     "collapse across all 100 training runs. This instability is an artifact of the small-foreground "
     "2D setting and does not arise in the paper's full-3D regime.")
para("**Experiments.** (1) main comparison of all four variants at full data; (2) a low-data sweep "
     "training each variant at 4, 8, 16 and all patients; (3) interpretability via spatial attention "
     "maps and the learned channel weights β. **Tools:** PyTorch, MONAI, nibabel, scikit-image, "
     "SciPy, matplotlib. Full experiments ran on an NVIDIA L40; the submitted Colab notebook "
     "reproduces the whole pipeline and defaults to a reduced proof-of-concept configuration.")

# ---- 4. Results ----
heading("4. Results and Analysis")
para("**Main comparison (full data, 5 seeds; pooled per-slice metrics over 1240 foreground "
     "slices).** All pairwise differences are statistically significant (paired Wilcoxon, p < 0.0001).")

rows = [
    ("Variant", "Dice", "Precision", "Recall"),
    ("U-Net (baseline)", "0.844", "0.865", "0.876"),
    ("Attention U-Net (original)", "0.865", "0.900", "0.892"),
    ("AG + CBAM (control)", "0.886", "0.907", "0.903"),
    ("Hybrid — ours", "0.844", "0.888", "0.863"),
]
table = doc.add_table(rows=0, cols=4)
table.style = "Light Grid Accent 1"
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, row in enumerate(rows):
    cells = table.add_row().cells
    for j, val in enumerate(row):
        cells[j].text = val
        for p in cells[j].paragraphs:
            for run in p.runs:
                run.font.size = Pt(9.5)
                if i == 0:
                    run.bold = True
caption("Table 1. Full-data segmentation results. AG+CBAM is best; the original Attention U-Net "
        "significantly improves over the plain U-Net; our hybrid ties the U-Net on Dice while "
        "trading recall for precision.")

para("**Channel attention significantly improves the paper's method.** Both attention variants "
     "beat the plain U-Net, and adding channel attention raises Dice further: the CBAM channel "
     "attention lifts the original Attention U-Net from 0.865 to **0.886** (p < 0.0001), the best "
     "result overall on every metric. This is the central positive finding — the Attention U-Net "
     "leaves channel selection unexploited, and adding it helps.")
para("**Our context-gated variant: a precision/recall trade-off.** At full data the hybrid ties "
     "the plain U-Net on Dice (0.844) and is significantly below the original Attention U-Net and "
     "CBAM. However, relative to the U-Net it *raises precision* (0.865 → 0.888) while *lowering "
     "recall* (0.876 → 0.863): the context-gated channel gate makes the model more conservative — "
     "predicting less spleen but more accurately — rather than more sensitive. This differs from the "
     "original paper's recall-boosting effect and indicates the extra channel gate is actively "
     "reshaping the decision, not merely adding capacity (it adds only 0.6% parameters over the "
     "spatial gate).")

figure("fig_low_data_curve.png", width_cm=12)
caption("Figure 1. Low-data regime: test Dice vs. training-set size (mean ± std, 5 seeds). Our "
        "hybrid (red) is the best variant at the smallest training size (4 patients), consistent "
        "with the hypothesis that context-gated channel selection helps most when data is scarce; "
        "the advantage does not persist at larger sizes.")

para("**Low-data regime.** Figure 1 tests our central hypothesis. At the smallest training set (4 "
     "patients) the hybrid is the top performer (Dice 0.696 vs. 0.670/0.656/0.655 for "
     "U-Net/Attention/CBAM), supporting the idea that a context-conditioned channel gate is most "
     "useful when per-image statistics are unreliable. The advantage is not monotonic — at 8 and 16 "
     "patients the hybrid is no longer ahead — so the evidence is suggestive of a low-data niche "
     "rather than a robust trend.")

figure("fig_channel_attn.png", width_cm=11)
caption("Figure 2. Learned context-gated channel weights β at three gated levels (sorted). β spans "
        "~0.1–1.0 (means 0.66–0.81, std ≈ 0.23), showing the gate performs meaningful, non-uniform "
        "channel selection rather than trivial pass-through.")

para("**Interpretability.** Figure 2 shows the learned channel weights β are far from uniform (they "
     "span roughly 0.1–1.0), confirming the gate learns genuine channel selection. Qualitative "
     "prediction overlays and spatial attention maps (in the repository) reproduce the paper's "
     "observation that gates localize the organ; the failure cases are dominated by thin organ tips "
     "and low-contrast boundary slices, where all variants struggle.")

# ---- 5. Conclusion ----
heading("5. Conclusion")
para("**Findings.** Starting from the hypothesis that context-gated channel attention would improve "
     "the Attention U-Net, we found a more nuanced result. Adding channel attention to the attention "
     "gate significantly improves spleen segmentation — a CBAM-style channel attention raises Dice "
     "from 0.865 to 0.886 (p < 0.0001) and is best on every metric. Our specific context-gated "
     "variant does not beat the original at full data; instead it shifts the model toward higher "
     "precision and lower recall, and its clearest benefit is the extreme low-data regime, where it "
     "is the best variant. The learned channel weights confirm the mechanism performs real, "
     "non-uniform channel selection.")
para("**Limitations.** A 2D proof-of-concept on a single organ at downsampled resolution, with a "
     "small dataset and high seed-to-seed variance for the gated variants; the required Dice+Focal "
     "loss adaptation; and, with 1240 paired samples, statistical significance that reflects small "
     "effect sizes. Results should be read as indicative, not definitive.")
para("**Future work.** Extend to full 3D volumes and multiple organs; add deep supervision (used in "
     "the original but omitted here); investigate why context-conditioning underperforms simple "
     "channel recalibration at full data while helping at low data; and test on thin, branching "
     "structures such as retinal vessels, where channel-specific selection may matter more.")

doc.save(OUT)
print("wrote", OUT)
